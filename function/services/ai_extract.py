"""
Read a document with OpenAI's vision models and return structured fields.

The call uses a strict JSON schema rather than asking for JSON in the prompt,
so the API itself rejects a malformed reply and there is nothing to parse
defensively on this side.

Scope: this module reads. It decides what a document appears to be and what is
written on it, and stops. Nothing here posts an invoice, touches payroll or
writes to AutoCount.
"""
import base64
import json
import mimetypes
import random
import subprocess
import tempfile
import time
import urllib.error
import urllib.request
from pathlib import Path

OPENAI_URL = "https://api.openai.com/v1/chat/completions"

# Classes a document can be filed under. `unclassified` is not in the list on
# purpose -- it is the state before a reading, not a possible reading.
DOC_CLASSES = ("sales", "purchase", "utility", "project_image", "salary", "other")

# The two firms whose books these are. Used by the classification rule below,
# and matched loosely because a letterhead is read off a photograph: spacing,
# "&" against "AND", and a trailing SDN BHD all vary.
OWN_COMPANIES = ("MANSON LIANG", "SENG CHONG")

# --- classification rule ---------------------------------------------------
#
# Which side of the transaction this business is on, decided by who issued the
# document rather than by whose name appears on it. A supplier invoice carries
# the supplier's letterhead and our name in the "to" field; an invoice we
# issued carries ours at the top and the customer's below.
#
# The first version of this had it backwards -- it filed a document as a
# purchase when the letterhead was ours -- and 105 of the first 138 uploads
# came out as sales when every one of them was something the company had
# bought. Naming both sides in the prompt, rather than "the letterhead", is
# what makes the question answerable from a photograph.
CLASSIFICATION_RULE = """\
- sales: a document THIS business issued to somebody else. The letterhead at
  the top is MANSON LIANG or SENG CHONG, and it is addressed to a customer.
- purchase: a document this business received for goods or work bought --
  materials, hardware, subcontracted labour, transport. The letterhead belongs
  to another business and MANSON LIANG or SENG CHONG appears as the recipient.
  If in doubt between sales and purchase, it is a purchase: this company
  receives far more documents than it issues.
- utility: a running cost rather than a purchase for a job -- water, electricity
  (TNB), fuel at a petrol station, internet and phone, rent, insurance. These
  are bills that arrive whether or not there is work on, which is what separates
  them from a purchase.
- salary: a payslip, salary voucher, EPF/KWSP statement, SOCSO/PERKESO form or
  any other document about paying staff, whoever the letterhead belongs to
- project_image: a photograph of work rather than a document -- cabinets, a
  site, an installation, materials delivered. No letterhead to read.
- other: anything that fits none of the above"""

SCHEMA = {
    "type": "object",
    "additionalProperties": False,
    "required": [
        "doc_class", "confidence", "issuer", "bill_to", "doc_no", "doc_date",
        "currency", "total", "summary", "people", "lines", "notes", "raw_text",
    ],
    "properties": {
        "doc_class": {"type": "string", "enum": list(DOC_CLASSES)},
        "confidence": {"type": "number"},
        "issuer": {"type": "string", "description": "Business on the letterhead. Empty if none."},
        "bill_to": {"type": "string", "description": "Who it is addressed to. Empty if none."},
        "doc_no": {"type": "string"},
        "doc_date": {"type": "string", "description": "YYYY-MM-DD, or empty if not shown."},
        "currency": {"type": "string"},
        "total": {"type": ["number", "null"]},
        "summary": {"type": "string", "description": "One sentence, for a list view."},
        "people": {
            "type": "array",
            "description": "Named individuals, for salary documents.",
            "items": {
                "type": "object",
                "additionalProperties": False,
                "required": ["name", "amount"],
                "properties": {
                    "name": {"type": "string"},
                    "amount": {"type": ["number", "null"]},
                },
            },
        },
        "lines": {
            "type": "array",
            "items": {
                "type": "object",
                "additionalProperties": False,
                "required": ["description", "qty", "unit_price", "amount"],
                "properties": {
                    "description": {"type": "string"},
                    "qty": {"type": ["number", "null"]},
                    "unit_price": {"type": ["number", "null"]},
                    "amount": {"type": ["number", "null"]},
                },
            },
        },
        "notes": {
            "type": "string",
            "description": "Anything unreadable, ambiguous or worth a human's attention.",
        },
        "raw_text": {
            "type": "string",
            "description": (
                "Every line of text on the document, verbatim, in reading order. "
                "Do not summarise, correct or reorder. Use [?] for a character "
                "you cannot make out."
            ),
        },
    },
}

PROMPT = f"""\
File this document for a Malaysian interior design and renovation business.

Classify it as exactly one of:
{CLASSIFICATION_RULE}

Then read what is on it.

- issuer is the business on the letterhead; bill_to is who the document is
  addressed to. Fill both when both are shown, and leave either empty rather
  than guessing.
- Dates are day-first in Malaysia: 03/04/2025 is 3 April 2025. Return
  YYYY-MM-DD, or an empty string if no date is printed.
- Numbers plain, no thousands separators. Use null, not 0, for an amount that
  is not shown -- a blank line and a zero mean different things on a payslip.
- Amounts written "RM" are MYR.
- For a salary document list each person named in `people`, with the amount
  against their name.
- Put anything you could not read, or read with difficulty, in `notes`. A note
  saying a figure was unclear is more useful than a confident wrong figure.
- confidence is your own, 0 to 1, for the classification.
- raw_text is a straight transcription of everything printed or written on the
  document, line by line. It is kept as the record of what was on the page, so
  transcribe rather than tidy: keep the original wording, spelling and order,
  and mark what you cannot read instead of guessing it.
"""


class ExtractionError(RuntimeError):
    """The document could not be read. The message is shown to the user."""


# A rate limit is not a failure, it is a queue asking to be slower. The account
# here allows 30,000 tokens a minute and a batch of scans at four at a time
# goes through that in seconds, so without this a large upload half-succeeds
# and leaves the rest marked failed for no reason anyone can act on.
RETRY_STATUSES = {429, 500, 502, 503, 504}
MAX_ATTEMPTS = 6


def _post_with_retry(body, *, api_key, timeout):
    request_body = json.dumps(body).encode()
    last = None
    for attempt in range(MAX_ATTEMPTS):
        request = urllib.request.Request(
            OPENAI_URL,
            data=request_body,
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        )
        try:
            with urllib.request.urlopen(request, timeout=timeout) as response:
                return json.load(response)
        except urllib.error.HTTPError as error:
            detail = ""
            try:
                detail = json.loads(error.read()).get("error", {}).get("message", "")
            except Exception:
                pass
            last = ExtractionError(f"OpenAI returned {error.code}: {detail or 'no detail'}")
            if error.code not in RETRY_STATUSES or attempt == MAX_ATTEMPTS - 1:
                raise last from error
            # Retry-After when the server says so, otherwise back off. The
            # jitter matters: four workers rate-limited by the same request
            # would otherwise wake together and hit the limit again.
            wait = _retry_after(error) or min(60, 2 ** attempt * 5)
            time.sleep(wait + random.uniform(0, 2))
        except (urllib.error.URLError, TimeoutError) as error:
            last = ExtractionError(f"Could not reach OpenAI: {error}")
            if attempt == MAX_ATTEMPTS - 1:
                raise last from error
            time.sleep(min(30, 2 ** attempt * 2) + random.uniform(0, 2))
    raise last or ExtractionError("OpenAI could not be reached.")


def _retry_after(error):
    raw = ""
    try:
        raw = error.headers.get("Retry-After") or ""
    except Exception:
        return None
    try:
        return max(0.0, min(120.0, float(raw)))
    except (TypeError, ValueError):
        return None


# --- what can be read, and how -------------------------------------------
#
# A file that cannot be read is still worth keeping; it is just never sent
# anywhere. Anything outside these sets is stored and skipped.
IMAGE_SUFFIXES = {".jpg", ".jpeg", ".png", ".webp", ".gif", ".tif", ".tiff", ".bmp"}
TEXT_DOC_SUFFIXES = {".xlsx", ".xlsm", ".docx"}
PDF_SUFFIXES = {".pdf"}
READABLE_SUFFIXES = IMAGE_SUFFIXES | TEXT_DOC_SUFFIXES | PDF_SUFFIXES

# Below this many characters a PDF's text layer is not worth trusting.
# Scanned pages routinely yield a scattering of characters from a stamp or a
# footer, which is worse than no text at all: it looks like a text document
# and reads like nothing. Anything thinner than this gets rendered and looked
# at instead.
MIN_PDF_TEXT_CHARS = 200

# Spreadsheets and long PDFs get truncated rather than refused. A document
# whose meaning is not in its first few thousand characters is not the kind of
# document this inbox is for.
MAX_TEXT_CHARS = 60_000


def is_readable(path):
    return Path(path).suffix.lower() in READABLE_SUFFIXES


def _pdf_text(path):
    """The PDF's own text layer, or "" if it has none worth reading."""
    try:
        out = subprocess.run(
            ["pdftotext", "-layout", "-q", str(path), "-"],
            check=True, capture_output=True, timeout=120,
        ).stdout.decode("utf-8", "replace")
    except (FileNotFoundError, subprocess.CalledProcessError, subprocess.TimeoutExpired):
        return ""
    return out if len(out.strip()) >= MIN_PDF_TEXT_CHARS else ""


def _xlsx_text(path):
    """Every sheet as tab-separated rows, blank rows and columns dropped."""
    try:
        from openpyxl import load_workbook
    except ImportError as error:
        raise ExtractionError("openpyxl is not installed; .xlsx cannot be read.") from error
    try:
        book = load_workbook(path, read_only=True, data_only=True)
    except Exception as error:  # openpyxl raises a zoo of exception types
        raise ExtractionError(f"Could not open the spreadsheet: {error}") from error

    chunks = []
    try:
        for sheet in book.worksheets:
            rows = []
            for row in sheet.iter_rows(values_only=True):
                cells = ["" if cell is None else str(cell) for cell in row]
                while cells and not cells[-1].strip():
                    cells.pop()
                if cells:
                    rows.append("\t".join(cells))
            if rows:
                chunks.append(f"--- sheet: {sheet.title} ---\n" + "\n".join(rows))
    finally:
        book.close()
    return "\n\n".join(chunks)


def _docx_text(path):
    """Paragraphs and table rows, in document order."""
    try:
        import docx
    except ImportError as error:
        raise ExtractionError("python-docx is not installed; .docx cannot be read.") from error
    try:
        document = docx.Document(str(path))
    except Exception as error:
        raise ExtractionError(f"Could not open the document: {error}") from error

    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n".join(parts)


def document_payload(path, *, page=1):
    """
    Decide how a file should be presented to the model.

    Text beats pixels wherever the file actually carries text: it is exact,
    cheaper, and does not depend on the model reading a photograph correctly.
    A PDF only falls back to rendering when its text layer is missing or too
    thin to trust.

    Returns {"kind": "text", "text": str} or
            {"kind": "image", "bytes": bytes, "mime": str}.
    """
    path = Path(path)
    suffix = path.suffix.lower()

    if suffix in TEXT_DOC_SUFFIXES:
        text = _xlsx_text(path) if suffix in {".xlsx", ".xlsm"} else _docx_text(path)
        if not text.strip():
            raise ExtractionError("The file has no readable content.")
        return {"kind": "text", "text": text[:MAX_TEXT_CHARS]}

    if suffix in PDF_SUFFIXES:
        text = _pdf_text(path)
        if text:
            return {"kind": "text", "text": text[:MAX_TEXT_CHARS]}
        data, mime = _render_pdf_page(path, page)
        return {"kind": "image", "bytes": data, "mime": mime}

    if suffix in IMAGE_SUFFIXES:
        return {"kind": "image", "bytes": path.read_bytes(),
                "mime": mimetypes.guess_type(path.name)[0] or "image/jpeg"}

    raise ExtractionError(f"{suffix or 'This file type'} is not something that can be read.")


def _render_pdf_page(path, page):
    """Render one page, for PDFs whose text layer is missing or too thin."""
    with tempfile.TemporaryDirectory() as tmp:
        stem = Path(tmp) / "page"
        try:
            subprocess.run(
                ["pdftoppm", "-r", "150", "-png", "-f", str(page), "-l", str(page),
                 str(path), str(stem)],
                check=True, capture_output=True, timeout=120,
            )
        except FileNotFoundError as error:
            raise ExtractionError("pdftoppm is not installed; PDFs cannot be read.") from error
        except subprocess.CalledProcessError as error:
            raise ExtractionError(
                f"Could not render page {page} of the PDF: "
                f"{error.stderr.decode('utf-8', 'replace')[:200]}"
            ) from error
        rendered = sorted(Path(tmp).glob("page*.png"))
        if not rendered:
            raise ExtractionError(f"Page {page} of the PDF produced no image.")
        return rendered[0].read_bytes(), "image/png"


def pdf_page_count(path):
    """Page count, or 1 for anything that is not a PDF or will not report."""
    if Path(path).suffix.lower() != ".pdf":
        return 1
    try:
        out = subprocess.run(["pdfinfo", str(path)], check=True, capture_output=True,
                             timeout=60).stdout.decode("utf-8", "replace")
    except (FileNotFoundError, subprocess.CalledProcessError, subprocess.TimeoutExpired):
        return 1
    for line in out.splitlines():
        if line.startswith("Pages:"):
            try:
                return max(1, int(line.split(":", 1)[1].strip()))
            except ValueError:
                break
    return 1


def _schema_without_transcription():
    """The same schema with the verbatim transcription dropped."""
    import copy
    schema = copy.deepcopy(SCHEMA)
    schema["required"] = [f for f in schema["required"] if f != "raw_text"]
    schema["properties"].pop("raw_text", None)
    return schema


def extract(path, *, api_key, model="gpt-4o", timeout=120, page=1):
    """
    Read one document. Returns (fields, source_text, usage).

    Raises ExtractionError for anything the caller should show a human: a
    refused key, a rate limit, an unreadable file. The caller records the
    message against the document rather than losing it in a log.
    """
    if not api_key:
        raise ExtractionError("No OPENAI_API_KEY is configured.")

    payload_in = document_payload(path, page=page)
    if payload_in["kind"] == "text":
        content = [{"type": "text",
                    "text": f"{PROMPT}\n\n--- document ---\n{payload_in['text']}"}]
    else:
        if not payload_in["bytes"]:
            raise ExtractionError("The file is empty.")
        encoded = base64.b64encode(payload_in["bytes"]).decode()
        content = [
            {"type": "text", "text": PROMPT},
            {"type": "image_url",
             "image_url": {"url": f"data:{payload_in['mime']};base64,{encoded}", "detail": "high"}},
        ]

    body = {
        "model": model,
        # A full transcription of a dense invoice plus the fields runs long;
        # at 4000 the JSON came back truncated and unparseable, which reads as
        # "the reply was not the expected shape" and hides the real cause.
        "max_tokens": 16000,
        # Reading a document is transcription, not composition. Anything above
        # zero invents plausible digits.
        "temperature": 0,
        "response_format": {
            "type": "json_schema",
            "json_schema": {"name": "document", "strict": True, "schema": SCHEMA},
        },
        "messages": [{"role": "user", "content": content}],
    }

    payload = _post_with_retry(body, api_key=api_key, timeout=timeout)

    # A dense page can be longer than any output limit. Rather than lose the
    # document, ask again for the fields alone -- knowing what the invoice says
    # without a transcription beats knowing nothing.
    if (payload.get("choices") or [{}])[0].get("finish_reason") == "length":
        body["response_format"]["json_schema"]["schema"] = _schema_without_transcription()
        payload = _post_with_retry(body, api_key=api_key, timeout=timeout)

    choice = (payload.get("choices") or [{}])[0]
    if choice.get("finish_reason") == "length":
        raise ExtractionError(
            "The document is too long to transcribe in one reply; the answer was cut off."
        )
    try:
        fields = json.loads(choice["message"]["content"])
    except (KeyError, IndexError, ValueError, TypeError) as error:
        raise ExtractionError("OpenAI's reply was not the expected shape.") from error

    # For a text document the source is what we sent; for a photograph the only
    # transcription that exists is the model's own.
    source_text = payload_in["text"] if payload_in["kind"] == "text" else str(fields.get("raw_text") or "")

    usage = payload.get("usage") or {}
    return fields, source_text, {
        "model": payload.get("model") or model,
        "tokens_in": int(usage.get("prompt_tokens") or 0),
        "tokens_out": int(usage.get("completion_tokens") or 0),
    }


def settle_class(proposed, issuer, bill_to):
    """
    Correct the model's classification where the parties settle it outright.

    The model reads a photograph and sometimes swaps the two sides of a form,
    which turns a supplier invoice into a sale. One direction is beyond
    argument: a document addressed to us was received by us, whatever the
    reading made of the letterhead.

    The mirror case -- our letterhead, addressed to somebody else -- is left
    alone, because that is also exactly what a genuine sales invoice looks like
    and there is no field that separates the two.
    """
    proposed = str(proposed or "other").strip() or "other"

    # These say what a document is about, not which side of a trade it is on.
    if proposed in ("salary", "project_image", "utility"):
        return proposed

    if looks_like_own_company(bill_to) and not looks_like_own_company(issuer):
        return "purchase"

    return proposed


def looks_like_own_company(name):
    """
    Whether a letterhead is one of the two firms.

    Compared with everything but letters and digits stripped, because the name
    is read off a photograph: "MANSON LIANG INTERIOR DESIGN & RENOVATION",
    "Manson Liang Interior Design and Renovation" and a version with the
    spacing lost all have to match.
    """
    squashed = "".join(ch for ch in str(name or "").upper() if ch.isalnum())
    if not squashed:
        return False
    return any(
        "".join(ch for ch in own if ch.isalnum()) in squashed
        for own in OWN_COMPANIES
    )
